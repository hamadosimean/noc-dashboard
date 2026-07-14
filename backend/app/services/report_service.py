import io

from sqlalchemy.orm import Session

from app.services import kpi_service


def build_report_data(db: Session, month: int, year: int) -> dict:
    summary = kpi_service.get_summary(db, month, year)
    localities = kpi_service.get_localities(db, month, year, limit=10)
    nodes = kpi_service.get_nodes(db, month, year, limit=10)
    recurrent = kpi_service.get_recurrent_nodes(db, month, year)
    return {
        "period": summary["period"],
        "kpi": summary["kpi"],
        "vs_previous_month": summary["vs_previous_month"],
        "top_localities": localities,
        "top_nodes": nodes,
        "recurrent_nodes": recurrent,
    }


def render_pdf(report: dict) -> bytes:
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            # En-tete bleu ANPTIC
            self.set_fill_color(0, 82, 155)
            self.rect(0, 0, 210, 20, "F")
            self.set_font("Helvetica", "B", 15)
            self.set_text_color(255, 255, 255)
            self.cell(0, 10, "Tableau de Bord NOC - Rapport Mensuel", border=0, ln=True, align="C")
            self.ln(10)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF()
    pdf.add_page()
    
    # Titre et periode
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 15, "Rapport d'Activite NOC", ln=True, align="C")
    
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, f"Periode : {report['period']['label']}", ln=True, align="C")
    pdf.ln(10)

    # Tableau KPI
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 10, "1. Indicateurs Cles de Performance (KPI)", ln=True)
    pdf.set_fill_color(240, 248, 255)
    pdf.set_draw_color(200, 200, 200)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(0, 0, 0)
    
    kpi = report["kpi"]
    labels = {
        "total_incidents": "Total incidents",
        "resolved": "Resolus",
        "open": "Ouverts",
        "resolution_rate_pct": "Taux de resolution (%)",
        "avg_mttr_minutes": "MTTR moyen (min)",
        "network_availability_pct": "Disponibilite reseau (%)",
        "critical_localities": "Localites critiques",
        "recurrent_nodes": "Noeuds recurrents",
        "off_hours_detected": "Incidents hors heures ouvrees",
    }
    
    for key, label in labels.items():
        # Ligne de tableau
        pdf.cell(110, 8, label, border=1, fill=True)
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(40, 8, str(kpi[key]), border=1, ln=True, align="C")
        pdf.set_font("Helvetica", "", 11)

    pdf.ln(10)

    # Tableau Top Localites
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 10, "2. Top Localites Affectees", ln=True)
    
    # En-tete du tableau
    pdf.set_fill_color(0, 82, 155)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(70, 8, "Localite (Region)", border=1, fill=True)
    pdf.cell(40, 8, "Incidents", border=1, fill=True, align="C")
    pdf.cell(40, 8, "Disponibilite (%)", border=1, ln=True, fill=True, align="C")
    
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 11)
    fill = False
    for loc in report["top_localities"]:
        pdf.set_fill_color(245, 245, 245) if fill else pdf.set_fill_color(255, 255, 255)
        pdf.cell(70, 8, f"{loc['locality']} ({loc['region']})", border=1, fill=fill)
        pdf.cell(40, 8, str(loc['total_incidents']), border=1, fill=fill, align="C")
        pdf.cell(40, 8, f"{loc['availability_pct']}%", border=1, ln=True, fill=fill, align="C")
        fill = not fill

    pdf.ln(10)

    # Noeuds Recurrents
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(0, 51, 102)
    pdf.cell(0, 10, "3. Noeuds Recurrents (>= 3 incidents)", ln=True)
    
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(0, 0, 0)
    if report["recurrent_nodes"]:
        for node in report["recurrent_nodes"]:
            pdf.cell(5, 7, "-", ln=False)
            pdf.set_font("Helvetica", "B", 11)
            pdf.cell(30, 7, node['code'], ln=False)
            pdf.set_font("Helvetica", "", 11)
            pdf.cell(0, 7, f"{node['name']} : {node['total_incidents']} incidents", ln=True)
    else:
        pdf.set_font("Helvetica", "I", 11)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 7, "Aucun noeud recurrent enregistre ce mois-ci.", ln=True)

    buffer = io.BytesIO()
    pdf.output(buffer)
    return buffer.getvalue()


KPI_LABELS = {
    "total_incidents": "Total incidents",
    "resolved": "Résolus",
    "open": "Ouverts",
    "resolution_rate_pct": "Taux de résolution (%)",
    "avg_mttr_minutes": "MTTR moyen (min)",
    "network_availability_pct": "Disponibilité réseau (%)",
    "critical_localities": "Localités critiques",
    "recurrent_nodes": "Nœuds récurrents",
    "off_hours_detected": "Incidents hors heures ouvrées",
}


def render_docx(report: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    
    # Titre stylisé
    title = doc.add_heading(level=0)
    title_run = title.add_run("Rapport Mensuel NOC — ANPTIC")
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"Période : {report['period']['label']}")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)
    sub_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tableau des KPIs
    doc.add_heading("1. Indicateurs clés de performance", level=1)
    kpi = report["kpi"]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Medium Shading 1 Accent 1"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Indicateur"
    hdr_cells[1].text = "Valeur"
    
    for key, label in KPI_LABELS.items():
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(kpi[key])

    doc.add_paragraph() # Espacement

    # Tableau Top Localités
    doc.add_heading("2. Top localités affectées", level=1)
    loc_table = doc.add_table(rows=1, cols=3)
    loc_table.style = "Medium Shading 1 Accent 1"
    
    hdr_cells = loc_table.rows[0].cells
    hdr_cells[0].text = "Localité (Région)"
    hdr_cells[1].text = "Incidents"
    hdr_cells[2].text = "Disponibilité (%)"

    for loc in report["top_localities"]:
        row_cells = loc_table.add_row().cells
        row_cells[0].text = f"{loc['locality']} ({loc['region']})"
        row_cells[1].text = str(loc['total_incidents'])
        row_cells[2].text = f"{loc['availability_pct']}%"

    doc.add_paragraph() # Espacement

    # Nœuds Récurrents
    doc.add_heading("3. Nœuds récurrents (≥ 3 incidents)", level=1)
    if report["recurrent_nodes"]:
        for node in report["recurrent_nodes"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{node['code']} ").bold = True
            p.add_run(f"({node['name']}) : {node['total_incidents']} incidents")
    else:
        p = doc.add_paragraph("Aucun nœud récurrent enregistré ce mois-ci.")
        p.italic = True

    # Pied de page
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Généré automatiquement par le tableau de bord NOC ANPTIC"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
